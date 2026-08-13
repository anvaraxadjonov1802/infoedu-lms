from pathlib import Path

from django.core.management.base import BaseCommand, CommandError
from docx import Document

from lms.importers.test_parser_v3 import inspect_document, parse_test_questions_document
from lms.management.commands.import_course_materials import discover_files


class Command(BaseCommand):
    help = 'Bitta test DOCX ichki formatini diagnostika uchun chiqaradi.'

    def add_arguments(self, parser):
        parser.add_argument('--root', default='import_materials')
        parser.add_argument('--topic', type=int, default=1)
        parser.add_argument('--limit', type=int, default=80)

    def handle(self, *args, **options):
        root = Path(options['root']).expanduser().resolve()
        if not root.exists():
            raise CommandError(f'Papka topilmadi: {root}')

        discovered = discover_files(root)
        path = discovered['tests'].get(options['topic'])
        if not path:
            raise CommandError(f"{options['topic']}-mavzu test DOCX topilmadi.")

        document = Document(str(path))
        self.stdout.write(f'Fayl: {path}')
        self.stdout.write(f'Paragraphs: {len(document.paragraphs)}, tables: {len(document.tables)}')
        self.stdout.write('--- DOCX ichki ko‘rinishi ---')
        for line in inspect_document(document, limit=options['limit']):
            self.stdout.write(line)

        parsed = parse_test_questions_document(document)
        self.stdout.write('--- Parser v3 natijasi ---')
        self.stdout.write(f'Aniqlangan savollar: {len(parsed)}')
        for item in parsed[:5]:
            self.stdout.write(
                f"{item['number']}. {item['question_text']} | "
                + ' | '.join(f"{o['id'].upper()}) {o['text']}" for o in item['options'])
                + f" | JAVOB={item['correct_answer'].upper()}"
            )
