import mimetypes
import re
from pathlib import Path

from django.core.files import File
from django.core.management.base import BaseCommand, CommandError
from django.db import transaction

from docx import Document

from lms.models import Course, Lesson, Module, Question, Test, TheoryContent, User
from lms.storage import upload_admin_file


QUESTION_RE = re.compile(r'^\s*(\d{1,3})\s*[\).:-]\s*(.+)$')
OPTION_RE = re.compile(r'^\s*([A-Ha-h])\s*[\).:-]\s*(.+)$')
ANSWER_RE = re.compile(
    r'^\s*(?:to[‘\'`]?g[‘\'`]?ri\s+javob|javob|answer)\s*[:\-]\s*([A-Ha-h])\b',
    re.IGNORECASE,
)
ANSWER_KEY_PAIR_RE = re.compile(r'(?<!\d)(\d{1,3})\s*[-:.)]?\s*([A-Ha-h])\b')


def clean_text(value):
    return re.sub(r'\s+', ' ', str(value or '')).strip()


def safe_read_docx(path):
    try:
        return Document(str(path))
    except Exception as exc:
        raise CommandError(f'DOCX ochilmadi: {path.name}: {exc}') from exc


def table_to_markdown(table):
    rows = []
    for row in table.rows:
        cells = [clean_text(cell.text).replace('|', '\\|') for cell in row.cells]
        if any(cells):
            rows.append(cells)
    if not rows:
        return ''
    width = max(len(row) for row in rows)
    rows = [row + [''] * (width - len(row)) for row in rows]
    header = rows[0]
    body = rows[1:]
    lines = [
        '| ' + ' | '.join(header) + ' |',
        '| ' + ' | '.join(['---'] * width) + ' |',
    ]
    lines.extend('| ' + ' | '.join(row) + ' |' for row in body)
    return '\n'.join(lines)


def docx_blocks(document):
    blocks = []
    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name or '').lower() if paragraph.style else ''
        is_heading = style.startswith('heading') or style in ('title', 'subtitle')
        blocks.append(('heading' if is_heading else 'text', text))

    for table in document.tables:
        md = table_to_markdown(table)
        if md:
            blocks.append(('table', md))
    return blocks


def topic_title_from_document(document, fallback):
    for kind, text in docx_blocks(document):
        if kind == 'heading' and 4 <= len(text) <= 220:
            return text
    for kind, text in docx_blocks(document):
        if kind == 'text' and 4 <= len(text) <= 180:
            return text
    return fallback


def sections_from_document(document):
    blocks = docx_blocks(document)
    if not blocks:
        return []

    sections = []
    current_title = 'Asosiy material'
    current = []

    def flush():
        nonlocal current
        content = '\n\n'.join(current).strip()
        if content:
            sections.append({'title': current_title, 'contentMarkdown': content})
        current = []

    saw_heading = False
    for kind, text in blocks:
        if kind == 'heading':
            saw_heading = True
            flush()
            current_title = text
        else:
            current.append(text)
    flush()

    if saw_heading and sections:
        return sections

    plain = [text for kind, text in blocks if kind != 'heading']
    if not plain:
        return sections

    grouped = []
    chunk = []
    for text in plain:
        chunk.append(text)
        if len(chunk) >= 6 or sum(len(x) for x in chunk) >= 2500:
            grouped.append(chunk)
            chunk = []
    if chunk:
        grouped.append(chunk)

    return [
        {
            'title': 'Asosiy material' if len(grouped) == 1 else f'{idx}-bo‘lim',
            'contentMarkdown': '\n\n'.join(items),
        }
        for idx, items in enumerate(grouped, start=1)
    ]


def summary_from_sections(sections):
    if not sections:
        return ''
    text = clean_text(sections[0].get('contentMarkdown', ''))
    return (text[:350] + '…') if len(text) > 350 else text


def topic_number_from_name(name):
    stem = Path(name).stem
    match = re.search(r'(?<!\d)(20|1[0-9]|[1-9])(?!\d)', stem)
    return int(match.group(1)) if match else None


def module_number_from_name(name):
    stem = Path(name).stem.upper()
    match = re.search(r'(?<!\d)([1-4])(?!\d)', stem)
    if match:
        return int(match.group(1))
    romans = {'I': 1, 'II': 2, 'III': 3, 'IV': 4}
    for roman in ('IV', 'III', 'II', 'I'):
        if re.search(rf'\b{roman}\b', stem):
            return romans[roman]
    return None


def discover_files(root):
    files = [p for p in root.rglob('*') if p.is_file() and not p.name.startswith('~$')]
    theory = {}
    tests = {}
    practical = {}
    independent = {}
    toc = None

    for path in files:
        name = path.name.lower()
        parent = path.parent.name.lower()
        if path.suffix.lower() != '.docx':
            continue
        if 'mundarija' in name:
            toc = path
            continue
        if 'test' in parent:
            number = topic_number_from_name(path.name)
            if number:
                tests[number] = path
            continue
        if 'amaliy' in parent or 'amaliy' in name:
            number = module_number_from_name(path.name)
            if number:
                practical[number] = path
            continue
        if 'mustaqil' in parent or 'mustaqil' in name:
            number = module_number_from_name(path.name)
            if number:
                independent[number] = path
            continue
        if 'nazariy' in parent or 'nazariya' in parent or 'mavzu' in name:
            number = topic_number_from_name(path.name)
            if number:
                theory[number] = path

    return {
        'theory': theory,
        'tests': tests,
        'practical': practical,
        'independent': independent,
        'toc': toc,
    }


def parse_toc_titles(path):
    if not path:
        return {}, {}
    document = safe_read_docx(path)
    topic_titles = {}
    module_titles = {}
    module_index = None

    for _, raw in docx_blocks(document):
        text = clean_text(raw)
        upper = text.upper()

        roman_match = re.search(r'\b(IV|III|II|I)\s*[-.]?\s*MODUL\b', upper)
        digit_match = re.search(r'\b([1-4])\s*[-.]?\s*MODUL\b', upper)
        if roman_match or digit_match:
            if digit_match:
                module_index = int(digit_match.group(1))
            else:
                module_index = {'I': 1, 'II': 2, 'III': 3, 'IV': 4}[roman_match.group(1)]
            if 3 <= len(text) <= 220:
                module_titles[module_index] = text
            continue

        match = re.match(r'^\s*(20|1[0-9]|[1-9])\s*[\).:-]?\s*(.+)$', text)
        if match:
            number = int(match.group(1))
            title = clean_text(match.group(2))
            title = re.sub(r'^(?:mavzu)\s*[:\-]?\s*', '', title, flags=re.IGNORECASE)
            if title:
                topic_titles[number] = title
    return topic_titles, module_titles


def parse_test_questions(path):
    document = safe_read_docx(path)
    lines = []

    for paragraph in document.paragraphs:
        text = clean_text(paragraph.text)
        if text:
            lines.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [clean_text(cell.text) for cell in row.cells if clean_text(cell.text)]
            if cells:
                lines.append(' | '.join(cells))

    answer_key = {}
    for line in lines:
        lower = line.lower()
        if 'javob' in lower or 'answer' in lower:
            for number, letter in ANSWER_KEY_PAIR_RE.findall(line):
                answer_key[int(number)] = letter.lower()

    parsed = []
    current = None

    def finish():
        nonlocal current
        if not current:
            return
        number = current['number']
        correct_letter = current.get('correct') or answer_key.get(number)
        options = current.get('options', [])
        if current.get('text') and len(options) >= 2 and correct_letter:
            ids = {item['id'] for item in options}
            if correct_letter in ids:
                parsed.append({
                    'number': number,
                    'question_text': current['text'],
                    'options': options,
                    'correct_answer': correct_letter,
                })
        current = None

    for line in lines:
        question_match = QUESTION_RE.match(line)
        if question_match:
            finish()
            current = {
                'number': int(question_match.group(1)),
                'text': clean_text(question_match.group(2)),
                'options': [],
                'correct': None,
            }
            continue

        if not current:
            continue

        answer_match = ANSWER_RE.match(line)
        if answer_match:
            current['correct'] = answer_match.group(1).lower()
            continue

        option_match = OPTION_RE.match(line)
        if option_match:
            option_id = option_match.group(1).lower()
            current['options'].append({
                'id': option_id,
                'text': clean_text(option_match.group(2)),
            })
            continue

        if not current['options'] and len(line) < 500:
            current['text'] = clean_text(f"{current['text']} {line}")

    finish()
    return parsed


def file_attachment(path, folder, upload):
    if not upload:
        return None
    content_type = mimetypes.guess_type(path.name)[0] or 'application/octet-stream'
    with path.open('rb') as raw:
        wrapped = File(raw, name=path.name)
        wrapped.content_type = content_type
        info = upload_admin_file(wrapped, folder=folder)
    return {
        'name': info['name'],
        'type': info['extension'],
        'size': info['sizeLabel'],
        'downloadUrl': info['url'],
    }


class Command(BaseCommand):
    help = 'Extract qilingan InfoEdu DOCX materiallarini LMS kursiga bulk import qiladi.'

    def add_arguments(self, parser):
        parser.add_argument('--root', default='import_materials', help='Extract qilingan materiallar papkasi.')
        parser.add_argument('--course-code', default='DAST-101')
        parser.add_argument('--course-title', default='Dasturlash')
        parser.add_argument('--teacher-email', default='admin@infoedu.uz')
        parser.add_argument('--publish', action='store_true', help='Importdan keyin kursni Published qiling.')
        parser.add_argument('--upload-files', action='store_true', help='Original DOCXlarni Supabase Storage ga yuklang.')
        parser.add_argument('--dry-run', action='store_true', help='Bazaga yozmasdan materiallarni tahlil qiling.')

    def handle(self, *args, **options):
        root = Path(options['root']).expanduser().resolve()
        if not root.exists() or not root.is_dir():
            raise CommandError(f'Papka topilmadi: {root}')

        discovered = discover_files(root)
        topic_titles, module_titles = parse_toc_titles(discovered['toc'])

        self.stdout.write(
            f"Topildi: nazariya={len(discovered['theory'])}, test={len(discovered['tests'])}, "
            f"amaliy={len(discovered['practical'])}, mustaqil={len(discovered['independent'])}"
        )

        missing_theory = [n for n in range(1, 21) if n not in discovered['theory']]
        missing_tests = [n for n in range(1, 21) if n not in discovered['tests']]
        if missing_theory:
            self.stdout.write(self.style.WARNING(
                f"Nazariya yetishmaydi: {', '.join(map(str, missing_theory))}-mavzu"
            ))
        if missing_tests:
            self.stdout.write(self.style.WARNING(
                f"Test yetishmaydi: {', '.join(map(str, missing_tests))}-mavzu"
            ))

        test_counts = {}
        for number, path in sorted(discovered['tests'].items()):
            try:
                test_counts[number] = len(parse_test_questions(path))
            except Exception as exc:
                test_counts[number] = 0
                self.stdout.write(self.style.WARNING(f'{number}-mavzu test o‘qilmadi: {exc}'))

        if options['dry_run']:
            for number in sorted(test_counts):
                self.stdout.write(f'{number}-mavzu test: {test_counts[number]} ta aniqlangan savol')
            self.stdout.write(self.style.SUCCESS('Dry-run yakunlandi. Bazaga o‘zgarish kiritilmadi.'))
            return

        teacher = User.objects.filter(email__iexact=options['teacher_email']).first()
        if not teacher:
            teacher = User.objects.filter(role__in=['admin', 'teacher']).order_by('id').first()
        if not teacher:
            raise CommandError('Teacher/Admin topilmadi. Avval o‘qituvchi yoki admin user yarating.')

        warnings = []

        with transaction.atomic():
            course, _ = Course.objects.update_or_create(
                code=options['course_code'],
                defaults={
                    'title': options['course_title'],
                    'description': 'InfoEdu bulk importer orqali yuklangan dasturlash kursi.',
                    'teacher': teacher,
                    'category': 'Dasturlash',
                    'level': 'Boshlang‘ich',
                    'status': 'published' if options['publish'] else 'draft',
                },
            )

            modules = {}
            for module_number in range(1, 5):
                module, _ = Module.objects.update_or_create(
                    course=course,
                    order=module_number,
                    defaults={
                        'title': module_titles.get(module_number, f'{module_number}-modul'),
                        'description': f'{module_number}-modul o‘quv materiallari.',
                    },
                )
                modules[module_number] = module

            for topic_number in range(1, 21):
                module_number = ((topic_number - 1) // 5) + 1
                within_module = ((topic_number - 1) % 5)
                module = modules[module_number]
                topic_title = topic_titles.get(topic_number, f'{topic_number}-mavzu')

                theory_path = discovered['theory'].get(topic_number)
                if theory_path:
                    document = safe_read_docx(theory_path)
                    sections = sections_from_document(document)
                    title = topic_titles.get(topic_number) or topic_title_from_document(document, topic_title)
                    lesson, _ = Lesson.objects.update_or_create(
                        module=module,
                        order=within_module * 2 + 1,
                        defaults={
                            'title': title,
                            'lesson_type': 'theory',
                            'duration_minutes': max(
                                5,
                                min(60, round(sum(len(s['contentMarkdown']) for s in sections) / 900)),
                            ),
                            'description': f'{topic_number}-mavzu nazariy materiali.',
                            'is_published': True,
                        },
                    )
                    attachments = []
                    try:
                        item = file_attachment(
                            theory_path,
                            f'courses/{course.code}/theory/topic-{topic_number}',
                            options['upload_files'],
                        )
                        if item:
                            attachments.append(item)
                    except Exception as exc:
                        warnings.append(f'{topic_number}-mavzu fayl upload: {exc}')

                    TheoryContent.objects.update_or_create(
                        lesson=lesson,
                        defaults={
                            'reading_time_minutes': max(5, lesson.duration_minutes),
                            'summary': summary_from_sections(sections),
                            'sections': sections,
                            'attachments': attachments,
                        },
                    )

                test_path = discovered['tests'].get(topic_number)
                if test_path:
                    lesson, _ = Lesson.objects.update_or_create(
                        module=module,
                        order=within_module * 2 + 2,
                        defaults={
                            'title': f'{topic_title} — Test',
                            'lesson_type': 'test',
                            'duration_minutes': 20,
                            'description': f'{topic_number}-mavzu bo‘yicha test.',
                            'is_published': True,
                        },
                    )
                    test, _ = Test.objects.update_or_create(
                        lesson=lesson,
                        defaults={
                            'title': f'{topic_title} testi',
                            'time_limit_minutes': 20,
                            'attempts_allowed': 3,
                            'passing_score_percent': 60,
                        },
                    )
                    questions = parse_test_questions(test_path)
                    if questions:
                        Question.objects.filter(test=test).delete()
                        Question.objects.bulk_create([
                            Question(
                                test=test,
                                question_text=item['question_text'],
                                question_type='single_choice',
                                options=item['options'],
                                correct_answer=item['correct_answer'],
                                explanation='',
                                points=1,
                                order=index,
                                topic=topic_title[:120],
                            )
                            for index, item in enumerate(questions, start=1)
                        ])
                    else:
                        warnings.append(
                            f'{topic_number}-mavzu testida avtomatik aniqlangan savol yo‘q; Test shell yaratildi.'
                        )

            for module_number, module in modules.items():
                practical_path = discovered['practical'].get(module_number)
                if practical_path:
                    document = safe_read_docx(practical_path)
                    sections = sections_from_document(document)
                    lesson, _ = Lesson.objects.update_or_create(
                        module=module,
                        order=11,
                        defaults={
                            'title': f'{module_number}-modul — Amaliy ish',
                            'lesson_type': 'practical',
                            'duration_minutes': 60,
                            'description': f'{module_number}-modul amaliy topshirig‘i.',
                            'is_published': True,
                        },
                    )
                    attachments = []
                    try:
                        item = file_attachment(
                            practical_path,
                            f'courses/{course.code}/practical/module-{module_number}',
                            options['upload_files'],
                        )
                        if item:
                            attachments.append(item)
                    except Exception as exc:
                        warnings.append(f'{module_number}-modul amaliy upload: {exc}')
                    TheoryContent.objects.update_or_create(
                        lesson=lesson,
                        defaults={
                            'reading_time_minutes': 15,
                            'summary': summary_from_sections(sections),
                            'sections': sections,
                            'attachments': attachments,
                        },
                    )

                independent_path = discovered['independent'].get(module_number)
                if independent_path:
                    document = safe_read_docx(independent_path)
                    sections = sections_from_document(document)
                    lesson, _ = Lesson.objects.update_or_create(
                        module=module,
                        order=12,
                        defaults={
                            'title': f'{module_number}-modul — Mustaqil ish',
                            'lesson_type': 'independent',
                            'duration_minutes': 90,
                            'description': f'{module_number}-modul mustaqil topshirig‘i.',
                            'is_published': True,
                        },
                    )
                    attachments = []
                    try:
                        item = file_attachment(
                            independent_path,
                            f'courses/{course.code}/independent/module-{module_number}',
                            options['upload_files'],
                        )
                        if item:
                            attachments.append(item)
                    except Exception as exc:
                        warnings.append(f'{module_number}-modul mustaqil upload: {exc}')
                    TheoryContent.objects.update_or_create(
                        lesson=lesson,
                        defaults={
                            'reading_time_minutes': 15,
                            'summary': summary_from_sections(sections),
                            'sections': sections,
                            'attachments': attachments,
                        },
                    )

        for warning in warnings:
            self.stdout.write(self.style.WARNING(warning))

        self.stdout.write(self.style.SUCCESS(
            f"Import tugadi: {course.code} — {course.title}. "
            f"Holat: {'Published' if options['publish'] else 'Draft'}."
        ))
        if not options['publish']:
            self.stdout.write(
                'Tekshiruvdan keyin admin paneldan kursni Published qiling yoki commandni --publish bilan qayta ishga tushiring.'
            )
