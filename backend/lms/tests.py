from pathlib import Path
from tempfile import TemporaryDirectory

from django.core.management import call_command
from django.test import TestCase
from docx import Document
from rest_framework.test import APIClient

from .management.commands.import_course_materials import parse_test_questions, topic_number_from_name
from .models import Course, Enrollment, User
from .storage import _headers, _safe_filename, format_file_size


class LMSApiSmokeTests(TestCase):
    @classmethod
    def setUpTestData(cls):
        call_command('seed_demo', verbosity=0)

    def setUp(self):
        self.client = APIClient()
        response = self.client.post('/api/auth/login/', {
            'email': 'anvar.axadjonov@tuit.uz',
            'password': 'InfoEdu2026!',
        }, format='json')
        self.assertEqual(response.status_code, 200)
        self.client.credentials(HTTP_AUTHORIZATION=f"Bearer {response.data['access']}")

    def test_bootstrap_masks_test_answers(self):
        response = self.client.get('/api/bootstrap/')
        self.assertEqual(response.status_code, 200)
        self.assertIn('courses', response.data)
        self.assertTrue(response.data['courses'])
        first_test = response.data['tests']['test-101']
        question = first_test['questions'][0]
        self.assertEqual(question.get('explanation'), '')
        self.assertNotIn('correctAnswer', question)

    def test_server_scores_attempt(self):
        response = self.client.post('/api/tests/test-101/submit/', {
            'answers': {
                'q-101-1': 'a',
                'q-101-2': 'b',
                'q-101-3': True,
            },
            'flaggedQuestionIds': [],
            'timeSpentSeconds': 95,
        }, format='json')
        self.assertEqual(response.status_code, 201)
        self.assertEqual(response.data['percentage'], 100)
        self.assertTrue(response.data['isPassed'])
        self.assertEqual(len(response.data['answerReviews']), 3)

    def test_new_student_is_auto_enrolled_in_all_published_courses(self):
        student = User.objects.create_user(
            username='auto-student',
            email='auto-student@example.com',
            password='SecurePass2026!',
            role='student',
        )
        published_ids = set(Course.objects.filter(status='published').values_list('id', flat=True))
        enrolled_ids = set(
            Enrollment.objects.filter(student=student).values_list('course_id', flat=True)
        )
        self.assertEqual(enrolled_ids, published_ids)

    def test_new_published_course_is_auto_enrolled_for_students(self):
        student = User.objects.create_user(
            username='future-student',
            email='future-student@example.com',
            password='SecurePass2026!',
            role='student',
        )
        teacher = User.objects.filter(role__in=['teacher', 'admin']).first()
        course = Course.objects.create(
            title='Yangi ochiq kurs',
            code='AUTO-ACCESS-COURSE',
            teacher=teacher,
            status='published',
        )
        self.assertTrue(Enrollment.objects.filter(student=student, course=course).exists())

    def test_student_cannot_access_admin_stats(self):
        response = self.client.get('/api/admin/stats/')
        self.assertEqual(response.status_code, 403)


class SupabaseStorageUnitTests(TestCase):
    def test_modern_secret_key_is_only_api_key_header(self):
        headers = _headers('sb_secret_example')
        self.assertEqual(headers['apikey'], 'sb_secret_example')
        self.assertNotIn('Authorization', headers)

    def test_legacy_service_role_key_can_be_bearer(self):
        headers = _headers('eyJlegacy-service-role')
        self.assertEqual(headers['apikey'], 'eyJlegacy-service-role')
        self.assertEqual(headers['Authorization'], 'Bearer eyJlegacy-service-role')

    def test_uploaded_filename_is_sanitized_and_unique(self):
        safe = _safe_filename('1-MAVZU O‘zbekcha (final).docx')
        self.assertTrue(safe.endswith('.docx'))
        self.assertNotIn(' ', safe)
        self.assertNotIn('‘', safe)

    def test_file_size_label(self):
        self.assertEqual(format_file_size(1024), '1.0 KB')


class BulkImporterUnitTests(TestCase):
    def test_topic_number_from_real_file_names(self):
        self.assertEqual(topic_number_from_name('1-MAVZU.docx'), 1)
        self.assertEqual(topic_number_from_name('10-mavzu.docx'), 10)
        self.assertEqual(topic_number_from_name('20-mavzu.docx'), 20)

    def test_single_choice_test_docx_is_parsed(self):
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / '1-mavzu.docx'
            document = Document()
            document.add_paragraph('1. Python nima?')
            document.add_paragraph('A) Operatsion tizim')
            document.add_paragraph('B) Dasturlash tili')
            document.add_paragraph('C) Brauzer')
            document.add_paragraph('Javob: B')
            document.save(path)

            questions = parse_test_questions(path)
            self.assertEqual(len(questions), 1)
            self.assertEqual(questions[0]['question_text'], 'Python nima?')
            self.assertEqual(questions[0]['correct_answer'], 'b')
            self.assertEqual(len(questions[0]['options']), 3)
