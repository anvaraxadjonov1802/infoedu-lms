from django.test import SimpleTestCase
from docx import Document

from lms.importers.test_parser_v3 import parse_test_questions_document


class RobustTestParserTests(SimpleTestCase):
    def test_explicit_question_format(self):
        document = Document()
        document.add_paragraph('1. Python nima?')
        document.add_paragraph('A) Dasturlash tili')
        document.add_paragraph('B) Operatsion tizim')
        document.add_paragraph('C) Brauzer')
        document.add_paragraph('D) Ma’lumotlar bazasi')
        document.add_paragraph('Javob: A')

        parsed = parse_test_questions_document(document)
        self.assertEqual(len(parsed), 1)
        self.assertEqual(parsed[0]['correct_answer'], 'a')
        self.assertEqual(parsed[0]['options'][0]['text'], 'Dasturlash tili')

    def test_table_question_format(self):
        document = Document()
        table = document.add_table(rows=2, cols=7)
        headers = ['№', 'Savol', 'A', 'B', 'C', 'D', 'Javob']
        values = ['1', 'Python nima?', 'Til', 'OS', 'Brauzer', 'DB', 'A']
        for idx, value in enumerate(headers):
            table.rows[0].cells[idx].text = value
        for idx, value in enumerate(values):
            table.rows[1].cells[idx].text = value

        parsed = parse_test_questions_document(document)
        self.assertEqual(len(parsed), 1)
        self.assertEqual(parsed[0]['question_text'], 'Python nima?')
        self.assertEqual(parsed[0]['correct_answer'], 'a')

    def test_table_can_use_marked_correct_option(self):
        document = Document()
        table = document.add_table(rows=1, cols=5)
        values = ['Python nima?', 'Til', 'OS', 'Brauzer', 'DB']
        for idx, value in enumerate(values):
            paragraph = table.rows[0].cells[idx].paragraphs[0]
            run = paragraph.add_run(value)
            if idx == 1:
                run.bold = True

        parsed = parse_test_questions_document(document)
        self.assertEqual(len(parsed), 1)
        self.assertEqual(parsed[0]['correct_answer'], 'a')

    def test_number_test_heading_with_separate_answer_key_table(self):
        document = Document()
        document.add_paragraph('1-MODUL. ALGORITMLASH ASOSLARI')
        document.add_paragraph('1-MAVZU. ALGORITM TUSHUNCHASI')
        document.add_paragraph('TEST TOPSHIRIQLARI')
        document.add_paragraph('1-test Algoritm nima?')
        document.add_paragraph('A) Kompyuter dasturi')
        document.add_paragraph('B) Buyruqlar va qadamlar ketma-ketligi')
        document.add_paragraph('C) Elektron qurilma')
        document.add_paragraph('D) Matn muharriri')
        document.add_paragraph('2-test Algoritmning bir bosqichi nima deyiladi?')
        document.add_paragraph('A) Natija')
        document.add_paragraph('B) Buyruq')
        document.add_paragraph('C) Qadam')
        document.add_paragraph('D) Dastur')

        table = document.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = 'Test №'
        table.rows[0].cells[1].text = 'To‘g‘ri javob'
        table.rows[1].cells[0].text = '1'
        table.rows[1].cells[1].text = 'B'
        table.rows[2].cells[0].text = '2'
        table.rows[2].cells[1].text = 'C'

        parsed = parse_test_questions_document(document)
        self.assertEqual(len(parsed), 2)
        self.assertEqual(parsed[0]['question_text'], 'Algoritm nima?')
        self.assertEqual(parsed[0]['correct_answer'], 'b')
        self.assertEqual(parsed[1]['correct_answer'], 'c')

    def test_repeated_test_answer_pairs_in_six_column_table(self):
        document = Document()
        for number in range(1, 7):
            document.add_paragraph(f'{number}-test Savol {number}?')
            document.add_paragraph('A) Birinchi')
            document.add_paragraph('B) Ikkinchi')
            document.add_paragraph('C) Uchinchi')
            document.add_paragraph('D) To‘rtinchi')

        table = document.add_table(rows=3, cols=6)
        headers = ['Test', 'Javob', 'Test', 'Javob', 'Test', 'Javob']
        for idx, value in enumerate(headers):
            table.rows[0].cells[idx].text = value
        values_1 = ['1', 'A', '3', 'C', '5', 'B']
        values_2 = ['2', 'B', '4', 'D', '6', 'A']
        for idx, value in enumerate(values_1):
            table.rows[1].cells[idx].text = value
        for idx, value in enumerate(values_2):
            table.rows[2].cells[idx].text = value

        parsed = parse_test_questions_document(document)
        self.assertEqual(len(parsed), 6)
        self.assertEqual([item['correct_answer'] for item in parsed], ['a', 'b', 'c', 'd', 'b', 'a'])

    def test_question_and_all_options_can_be_in_one_paragraph(self):
        document = Document()
        document.add_paragraph(
            '1. Yakuniy loyiha yaratishdan asosiy maqsad nima? '
            'A) Bilimlarni amalda qo‘llash B) Faqat rasm chizish '
            'C) Kompyuter qismlari D) Internetdan foydalanish'
        )
        document.add_paragraph(
            '2. Loyiha bosqichlarining birinchisi qaysi? '
            'A) Sinov B) Reja C) Taqdimot D) Dasturlash'
        )
        table = document.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = '№'
        table.rows[0].cells[1].text = 'To‘g‘ri javob'
        table.rows[1].cells[0].text = '1'
        table.rows[1].cells[1].text = 'A'
        table.rows[2].cells[0].text = '2'
        table.rows[2].cells[1].text = 'B'

        parsed = parse_test_questions_document(document)
        self.assertEqual(len(parsed), 2)
        self.assertEqual(parsed[0]['question_text'], 'Yakuniy loyiha yaratishdan asosiy maqsad nima?')
        self.assertEqual(parsed[0]['correct_answer'], 'a')
        self.assertEqual(parsed[1]['options'][1]['text'], 'Reja')
